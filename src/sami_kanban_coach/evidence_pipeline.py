"""Evidence pipeline for local-model-backed mailbox evidence recall.

Layered search → attachment extraction → parsing → manifest → model input.

Safety: read-only by design. No mailbox or Kanban mutations.
"""

from __future__ import annotations

import hashlib
import json
import os
import re
import shutil
import uuid
from datetime import datetime, timezone
from enum import Enum
from pathlib import Path
from typing import Any

# ── Search status enum ─────────────────────────────────────────────────────
class SearchStatus(str, Enum):
    NOT_SEARCHED = "not_searched"
    NO_MATCH = "no_match"
    SUBJECT_MATCH_ONLY = "subject_match_only"
    BODY_CONTEXT_FOUND = "body_context_found"
    CONVERSATION_THREAD_FOUND = "conversation_thread_found"
    ATTACHMENT_METADATA_FOUND = "attachment_metadata_found"
    ATTACHMENT_EVIDENCE_FOUND = "attachment_evidence_found"
    SENT_ITEMS_EVIDENCE_FOUND = "sent_items_evidence_found"
    PARTIAL_RELATED_CONTEXT_ONLY = "partial_related_context_only"
    CONFLICTING_EVIDENCE_FOUND = "conflicting_evidence_found"
    ERROR_READ_ONLY_FAILURE = "error_read_only_failure"


class EvidenceStrength(str, Enum):
    STRONG = "strong"          # spreadsheet/doc with target terms
    MODERATE = "moderate"      # email body with SRV/REQ/IP terms
    WEAK = "weak"              # subject match only, related thread
    INCONCLUSIVE = "inconclusive"
    NONE = "none"


# ── unsafe types ───────────────────────────────────────────────────────────
_UNSAFE_EXTENSIONS = frozenset({
    ".exe", ".com", ".bat", ".cmd", ".ps1", ".vbs", ".js", ".jar",
    ".msi", ".scr", ".pif", ".reg", ".wsf", ".py", ".sh",
})


# ── Config helpers ─────────────────────────────────────────────────────────
def _load_settings(settings_path: Path) -> dict:
    with open(settings_path, encoding="utf-8-sig") as f:
        return json.load(f)


def _save_settings(settings_path: Path, data: dict) -> None:
    with open(settings_path, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2)


def _enable_mailbox(settings: dict, settings_path: Path) -> dict:
    old = {
        "enabled": settings.get("mailbox_search_enabled", False),
        "days": settings.get("mailbox_search_recent_days", 180),
    }
    if not settings.get("mailbox_search_enabled", False):
        settings["mailbox_search_enabled"] = True
        settings["mailbox_search_recent_days"] = 365
        _save_settings(settings_path, settings)
    return old


def _restore_mailbox(settings: dict, settings_path: Path, old: dict) -> None:
    settings["mailbox_search_enabled"] = bool(old.get("enabled", False))
    settings["mailbox_search_recent_days"] = int(old.get("days", 180))
    _save_settings(settings_path, settings)


# ── Sanitize helper ────────────────────────────────────────────────────────
def _sanitize(name: str) -> str:
    return re.sub(r"[<>:\"/\\|?*]", "_", re.sub(r"\s+", "_", name)).strip("._")


# ── Tesseract / OCR ────────────────────────────────────────────────────────
def detect_tesseract() -> dict:
    """Detect tesseract path and version. Public version of _detect_tesseract."""
    candidates = [
        r"C:\Tools\Tesseract-OCR\tesseract.exe",
    ]
    # Check PATH
    import shutil as _shutil
    path_tess = _shutil.which("tesseract")
    if path_tess:
        candidates.insert(0, path_tess)

    for path in candidates:
        p = Path(path)
        if p.exists():
            try:
                import subprocess
                out = subprocess.check_output([str(p), "--version"], stderr=subprocess.STDOUT, timeout=5)
                ver = out.decode().split("\n")[0].strip()
                # Check eng traineddata
                tessdata = p.parent / "tessdata"
                if not tessdata.exists():
                    tessdata = Path(re.sub(r"/tesseract\.exe$", "/tessdata", str(p), flags=re.I))
                eng_ok = (tessdata / "eng.traineddata").exists()
                return {
                    "available": True,
                    "path": str(p),
                    "version": ver,
                    "tessdata_path": str(tessdata),
                    "eng_available": eng_ok,
                }
            except Exception as e:
                return {"available": False, "error": str(e)}
    return {"available": False, "error": "tesseract not found"}


def _ocr_image(image_path: Path, tesseract_info: dict) -> dict:
    """OCR a single image. Returns structured result."""
    if not tesseract_info.get("available"):
        return {"source": str(image_path), "ocr_available": False, "text": "", "error": "tesseract unavailable"}
    try:
        import pytesseract
        from PIL import Image
        pytesseract.pytesseract.tesseract_cmd = tesseract_info["path"]
        os.environ["TESSDATA_PREFIX"] = tesseract_info.get("tessdata_path", "")
        img = Image.open(str(image_path)).convert("RGB")
        text = pytesseract.image_to_string(img, lang="eng").strip()
        # Classify
        if not text:
            classification = "blank_or_illegible"
        elif len(text) < 20:
            classification = "logo_or_signature"
        elif len(text) < 100:
            classification = "screenshot_or_label"
        else:
            classification = "document_scan"
        return {
            "source": str(image_path),
            "ocr_available": True,
            "text": text[:2000],
            "char_count": len(text),
            "classification": classification,
        }
    except Exception as e:
        return {"source": str(image_path), "ocr_available": False, "text": "", "error": str(e)}


# ── Evidence target terms ──────────────────────────────────────────────────
EVIDENCE_TERMS = [
    # SRV/REQ
    r"SRV-\d{5,7}", r"REQ\d{6,7}", r"SR#\s*\d{5,8}",
    # IP addresses
    r"\b\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}\b",
    # People
    r"\bSiva\b", r"\bFaraz\b", r"\bJason\b", r"Daniel Schroeder",
    r"Nicholas Thong", r"Kuruvilla",
    # Systems
    r"\bNEC\b", r"\bOCIO\b", r"\bIPESC\b", r"NTGMIPRDG",
    r"NT DCDD", r"DCDD",
    # Domains
    r"\bUltraRad\b", r"\bRAH\b", r"NT Health", r"\bStroke\b",
    r"\bVPN\b", r"\bFirewall\b", r"\bDICOM\b", r"\bPACS\b",
    r"SA PACS", r"Stroke PACS",
    r"DHSA Firewall Technical Requirements",
    r"\bSIA\b", r"Technical Requirements", r"change request",
    r"IPESC VPN", r"\bTHS\b", r"DoH firewall",
    r"DICOM Connectivity",
    # AE titles
    r"\bAE\b", r"AE Title",
    # Ports
    r"\bTCP \d{2,5}\b", r"\bport \d{2,5}\b",
    # Vendors
    r"\bNEC\b", r"\bTelstra\b", r"\bSiemens\b",
    # NT IP ranges
    r"10\.2\.\d{1,3}\.\d{1,3}",
    # SA IP ranges
    r"10\.\d{1,3}\.\d{1,3}\.\d{1,3}",
]


def _compile_term_patterns() -> list[re.Pattern]:
    patterns = []
    for t in EVIDENCE_TERMS:
        try:
            patterns.append(re.compile(t, re.IGNORECASE))
        except re.error:
            pass
    return patterns


_TERM_PATTERNS = _compile_term_patterns()


def search_text_for_terms(text: str) -> dict[str, list[str]]:
    """Search text for all evidence terms. Returns {term_category: [matches]}."""
    if not text:
        return {"srv": [], "req": [], "ips": [], "people": [], "systems": [],
                "domains": [], "ports": [], "vendors": [], "all": []}
    matches = {"srv": [], "req": [], "ips": [], "people": [], "systems": [],
               "domains": [], "ports": [], "vendors": [], "all": []}
    found_set = set()

    for p in _TERM_PATTERNS:
        for m in p.finditer(text):
            val = m.group(0)
            if val in found_set:
                continue
            found_set.add(val)
            matches["all"].append(val)
            pat_str = p.pattern
            if pat_str.startswith(r"SRV-"):
                matches["srv"].append(val)
            elif pat_str.startswith(r"REQ"):
                matches["req"].append(val)
            elif pat_str.startswith(r"SR#"):
                matches["req"].append(val)
            elif r"\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}" in pat_str:
                matches["ips"].append(val)
            elif any(n in pat_str for n in ["Siva", "Faraz", "Jason", "Schroeder", "Thong", "Kuruvilla"]):
                matches["people"].append(val)
            elif any(n in pat_str for n in ["NEC", "OCIO", "IPESC", "NTGMIPRDG", "DCDD"]):
                matches["systems"].append(val)
            elif any(n in pat_str for n in ["TCP", "port"]):
                matches["ports"].append(val)
            elif any(n in pat_str for n in ["Telstra", "Siemens"]):
                matches["vendors"].append(val)
            else:
                matches["domains"].append(val)
    return matches


# ── Outlook connection ─────────────────────────────────────────────────────
def _connect_outlook():
    """Connect to Outlook COM. Returns (app, ns, inbox, sent)."""
    import win32com.client
    app = win32com.client.Dispatch("Outlook.Application")
    ns = app.GetNamespace("MAPI")
    inbox = ns.GetDefaultFolder(6)   # olFolderInbox
    sent = ns.GetDefaultFolder(5)    # olFolderSentMail
    return app, ns, inbox, sent


def _get_folder_info(folder) -> dict:
    info = {}
    try: info["folderPath"] = str(folder.FolderPath)
    except: info["folderPath"] = "?"
    try: info["folderName"] = str(folder.Name)
    except: info["folderName"] = "?"
    try:
        st = folder.Store
        info["storeDisplayName"] = str(getattr(st, "DisplayName", "?"))
    except: info["storeDisplayName"] = "?"
    try: info["totalItemCount"] = int(folder.Items.Count)
    except: info["totalItemCount"] = 0
    try: info["entryID"] = str(folder.EntryID or "")[:32]
    except: info["entryID"] = ""
    return info


def _process_item(item, source_folder: str, tag: str) -> dict | None:
    """Extract fields from a mail item into a structured dict."""
    subj = sender = sndr_email = body = entry_id = conv_id = conv_topic = ""
    recv = None
    try: subj = str(item.Subject or "")
    except: pass
    try: recv = item.ReceivedTime
    except: pass
    try: sender = str(item.SenderName or "")
    except: pass
    try: sndr_email = str(item.SenderEmailAddress or "")
    except: pass
    try: body = str(item.Body or "")
    except: pass
    try: entry_id = str(item.EntryID or "")
    except: pass
    try: conv_id = str(item.ConversationID or "")
    except: pass
    try: conv_topic = str(item.ConversationTopic or "")
    except: pass

    recv_str = recv.isoformat() if recv else ""
    dedup_raw = f"{entry_id}:{subj}:{sndr_email}:{recv_str}"
    dedup_hash = hashlib.sha256(dedup_raw.encode()).hexdigest()[:24]
    email_key = f"{tag}:{dedup_hash}"

    # Extract to/cc
    to_list = []; cc_list = []
    try:
        for r in item.Recipients:
            try:
                addr = str(r.Address or "")
                name = str(r.Name or "")
                entry = f"{name} <{addr}>" if addr else name
                tp = r.Type
                if tp == 1:
                    to_list.append(entry)
                elif tp == 2:
                    cc_list.append(entry)
            except: pass
    except: pass

    return {
        "email_key": email_key,
        "subject": subj[:200],
        "sender": sender,
        "sender_email": sndr_email,
        "to_recipients": to_list,
        "cc_recipients": cc_list,
        "date": recv_str,
        "body": body,
        "body_preview": body[:500],
        "entry_id": entry_id[:40],
        "conversation_id": conv_id,
        "conversation_topic": conv_topic,
        "source_folder": source_folder,
        "_item": item,
        "_raw_received": recv,
    }


# ── Layered search ─────────────────────────────────────────────────────────
def search_by_subject(folder, subject_patterns: list[str], tag: str, source_folder: str) -> list[dict]:
    """Search a folder by subject patterns using DASL Restrict."""
    results = []
    seen_keys = set()
    for pat in subject_patterns:
        escaped = pat.replace("'", "''")
        dasl = f'@SQL="urn:schemas:httpmail:subject" LIKE \'%{escaped}%\''
        try:
            items = folder.Items.Restrict(dasl)
            items.Sort("[ReceivedTime]", False)
            for item in items:
                processed = _process_item(item, source_folder, tag)
                if processed and processed["email_key"] not in seen_keys:
                    seen_keys.add(processed["email_key"])
                    results.append(processed)
        except Exception:
            pass
    return results


def search_by_conversation(folder, conv_id: str, exclude_keys: set, tag: str, source_folder: str) -> list[dict]:
    """Expand by ConversationID."""
    results = []
    if not conv_id or len(conv_id) < 4:
        return results
    escaped = conv_id.replace("'", "''")
    dasl = f'@SQL="urn:schemas:httpmail:conversationid" = \'{escaped}\''
    try:
        items = folder.Items.Restrict(dasl)
        for item in items:
            processed = _process_item(item, source_folder, tag)
            if processed and processed["email_key"] not in exclude_keys:
                results.append(processed)
    except Exception:
        pass
    return results


def search_by_body_keywords(folder, keywords: list[str], max_results: int,
                            tag: str, source_folder: str) -> list[dict]:
    """Chunked month-by-month scan with keyword body search."""
    results = []
    seen_keys = set()
    today = datetime.now(timezone.utc)
    current = today.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
    start_boundary = today - timedelta(days=365)
    patterns = []
    for kw in keywords:
        if len(kw) >= 3:
            try: patterns.append(re.compile(re.escape(kw), re.IGNORECASE))
            except: pass
    if not patterns:
        return results

    chunk_start = max(current, start_boundary)
    while chunk_start >= start_boundary and len(results) < max_results:
        chunk_end = (chunk_start + timedelta(days=32)).replace(day=1) - timedelta(days=1)
        if chunk_end > today:
            chunk_end = today
        start_str = chunk_start.strftime("%m/%d/%Y")
        end_str = chunk_end.strftime("%m/%d/%Y")
        month_filter = (
            f'@SQL="urn:schemas:httpmail:datereceived" >= \'{start_str}\''
            f' AND "urn:schemas:httpmail:datereceived" <= \'{end_str}\''
        )
        try:
            items = folder.Items.Restrict(month_filter)
            for item in items:
                if len(results) >= max_results:
                    break
                processed = _process_item(item, source_folder, tag)
                if not processed or processed["email_key"] in seen_keys:
                    continue
                # Check date cutoff
                rd = processed["_raw_received"]
                if rd is not None:
                    try:
                        if hasattr(rd, "tzinfo") and rd.tzinfo is None:
                            rd = rd.replace(tzinfo=timezone.utc)
                        if rd < start_boundary:
                            break
                    except: pass
                # Keyword match
                body_text = processed.get("body", "")
                subj_text = processed.get("subject", "")
                full_text = f"{subj_text} {body_text}"
                for p in patterns:
                    if p.search(full_text):
                        seen_keys.add(processed["email_key"])
                        results.append(processed)
                        break
        except Exception:
            pass
        chunk_start = (chunk_start - timedelta(days=1)).replace(day=1)
    return results[:max_results]


def search_attachments_in_folder(folder, keyword: str, tag: str, source_folder: str) -> list[dict]:
    """Search for emails with attachments matching a filename keyword."""
    results = []
    seen_keys = set()
    kw_lower = keyword.lower()
    try:
        for item in folder.Items:
            processed = _process_item(item, source_folder, tag)
            if not processed or processed["email_key"] in seen_keys:
                continue
            try:
                for att in item.Attachments:
                    name = str(getattr(att, "DisplayName", "") or getattr(att, "FileName", "") or "")
                    if kw_lower in name.lower():
                        seen_keys.add(processed["email_key"])
                        results.append(processed)
                        break
            except: pass
    except: pass
    return results


# ── Attachment extraction ──────────────────────────────────────────────────
def extract_attachments(email: dict, export_dir: Path) -> list[dict]:
    """Extract attachments from a single email to export_dir. Returns list of att dicts."""
    att_list = []
    item = email.get("_item")
    if not item:
        return att_list
    try:
        for i, att in enumerate(item.Attachments, start=1):
            try:
                name = str(getattr(att, "DisplayName", "") or getattr(att, "FileName", "") or f"att_{i}")
                att_size = int(getattr(att, "Size", 0) or 0)
                ext = Path(name).suffix.lower() or ".bin"
                if ext in _UNSAFE_EXTENSIONS:
                    att_list.append({
                        "original_filename": name, "extension": ext, "size": att_size,
                        "saved": False, "sha256": "", "parse_status": "unsafe_type",
                        "email_subject": email.get("subject", ""),
                        "sender": email.get("sender", ""),
                        "date": email.get("date", ""),
                        "email_key": email.get("email_key", ""),
                    })
                    continue
                safe_name = _sanitize(name)
                ts = datetime.now().strftime("%Y%m%d_%H%M%S")
                slug = _sanitize(email.get("subject", "nosub")[:30])
                save_name = f"{ts}_{slug}_{i}_{safe_name}"
                save_path = export_dir / save_name

                # Save via COM
                tmp = Path(os.environ.get("TEMP", "/tmp")) / f"ep_att_{uuid.uuid4().hex[:8]}{ext}"
                att.SaveAsFile(str(tmp))
                sha256 = hashlib.sha256(tmp.read_bytes()).hexdigest()
                shutil.copy2(str(tmp), str(save_path))
                try: tmp.unlink()
                except: pass

                att_list.append({
                    "original_filename": name,
                    "sanitized_filename": save_name,
                    "saved_path": str(save_path),
                    "extension": ext,
                    "size": att_size,
                    "sha256": sha256,
                    "saved": True,
                    "parse_status": "pending",
                    "email_subject": email.get("subject", ""),
                    "sender": email.get("sender", ""),
                    "date": email.get("date", ""),
                    "email_key": email.get("email_key", ""),
                })
            except Exception as e:
                att_list.append({
                    "original_filename": name if 'name' in dir() else f"att_{i}",
                    "extension": ".bin", "size": 0,
                    "saved": False, "sha256": "", "parse_status": f"extract_error: {e}",
                })
    except Exception as e:
        pass
    return att_list


# ── Attachment parsing ─────────────────────────────────────────────────────
def parse_attachment(att: dict, tesseract_info: dict | None = None) -> dict:
    """Parse a saved attachment. Returns att dict with parsed content."""
    if not att.get("saved") or not att.get("saved_path"):
        att["parse_status"] = att.get("parse_status", "not_saved")
        att["text_content"] = ""
        att["structured"] = {}
        att["term_matches"] = search_text_for_terms("")
        return att

    fp = Path(att["saved_path"])
    if not fp.exists():
        att["parse_status"] = "file_missing"
        return att

    ext = att.get("extension", "").lower()
    text = ""
    structured = {}

    try:
        if ext in (".xlsx", ".xlsm", ".xls"):
            import openpyxl
            wb = openpyxl.load_workbook(fp, read_only=True, data_only=True)
            sheet_info = {}
            for ws in wb.worksheets:
                rows_data = []
                for row in ws.iter_rows(values_only=True):
                    row_vals = [str(c) if c is not None else "" for c in row]
                    rows_data.append(row_vals)
                sheet_info[ws.title] = {
                    "row_count": len(rows_data),
                    "non_empty_rows": sum(1 for r in rows_data if any(c.strip() for c in r)),
                }
                # Concatenate for text search
                for r in rows_data:
                    text += " | ".join(c for c in r if c.strip()) + "\n"
            structured["sheets"] = sheet_info
            wb.close()
            att["parse_status"] = "parsed"

        elif ext == ".csv":
            text = fp.read_text(encoding="utf-8", errors="replace")
            att["parse_status"] = "parsed"

        elif ext == ".txt":
            text = fp.read_text(encoding="utf-8", errors="replace")
            att["parse_status"] = "parsed"

        elif ext == ".docx":
            import docx
            doc = docx.Document(fp)
            # Paragraphs
            for p in doc.paragraphs:
                text += p.text + "\n"
            # Tables
            for table in doc.tables:
                for row in table.rows:
                    text += " | ".join(cell.text for cell in row.cells) + "\n"
                text += "---\n"
            att["parse_status"] = "parsed"

        elif ext == ".pdf":
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(fp)
                pages = []
                for i, page in enumerate(reader.pages):
                    pt = page.extract_text() or ""
                    pages.append({"page": i + 1, "text": pt[:500]})
                    text += pt + "\n"
                structured["pages"] = pages
                att["parse_status"] = "parsed"
            except Exception:
                try:
                    from pdfminer.high_level import extract_text as pdf_extract
                    text = pdf_extract(str(fp))
                    att["parse_status"] = "parsed"
                except ImportError:
                    att["parse_status"] = "parse_error_pdf"
                    text = ""

        elif ext in (".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tif", ".tiff"):
            ocr_result = _ocr_image(fp, tesseract_info or {})
            text = ocr_result.get("text", "")
            structured["ocr"] = {
                "available": ocr_result.get("ocr_available", False),
                "char_count": ocr_result.get("char_count", 0),
                "classification": ocr_result.get("classification", "unknown"),
            }
            att["parse_status"] = "ocr_parsed" if text else "no_text"

        elif ext == ".msg":
            att["parse_status"] = "msg_not_parsed"

        else:
            att["parse_status"] = f"unsupported_format: {ext}"

    except Exception as e:
        att["parse_status"] = f"parse_error: {e}"
        text = ""

    att["text_content"] = text[:50000]
    att["structured"] = structured
    att["term_matches"] = search_text_for_terms(text)
    return att


# ── Evidence manifest ──────────────────────────────────────────────────────
def build_evidence_manifest(
    run_id: str,
    search_results: dict,
    attachments: list[dict],
    inbox_info: dict,
    sent_info: dict,
    settings: dict,
) -> dict:
    """Build the evidence_manifest.json."""
    manifest = {
        "run_id": run_id,
        "timestamp": datetime.now().isoformat(),
        "pipeline_version": "1.0",
        "searched_folders": {
            "inbox": inbox_info,
            "sent_items": sent_info,
        },
        "search_results": {
            "status": search_results.get("status", SearchStatus.NOT_SEARCHED.value),
            "inbox_messages_matched": len(search_results.get("inbox_matches", [])),
            "sent_messages_matched": len(search_results.get("sent_matches", [])),
            "conversation_ids_found": search_results.get("conv_ids", []),
            "body_keyword_search_used": search_results.get("body_keyword_used", False),
            "attachment_search_used": search_results.get("att_search_used", False),
            "total_items_scanned_inbox": search_results.get("inbox_scanned", 0),
            "total_items_scanned_sent": search_results.get("sent_scanned", 0),
        },
        "attachments": {
            "total_extracted": len(attachments),
            "saved_count": sum(1 for a in attachments if a.get("saved")),
            "parsed_count": sum(1 for a in attachments if a.get("parse_status", "").startswith("parsed")),
            "ocr_parsed_count": sum(1 for a in attachments if a.get("parse_status") == "ocr_parsed"),
            "unsafe_skipped": sum(1 for a in attachments if a.get("parse_status") == "unsafe_type"),
            "error_count": sum(1 for a in attachments if a.get("parse_status", "").startswith("parse_error") or a.get("parse_status", "").startswith("extract_error")),
        },
        "ocr": detect_tesseract(),
        "config_snapshot": {
            "mailbox_search_enabled": settings.get("mailbox_search_enabled", False),
            "mailbox_search_recent_days": settings.get("mailbox_search_recent_days", 180),
            "allow_kanban_apply": settings.get("allow_kanban_apply", False),
            "local_kanban_apply_enabled": settings.get("local_kanban_apply_enabled", False),
        },
        "safety": {
            "mailboxMutated": False,
            "kanbanWritePerformed": False,
            "teamEsmiWritePerformed": False,
        },
        "errors": [],
        "warnings": [],
    }
    return manifest


# ── Search status classifier ──────────────────────────────────────────────
def classify_search(outbox_matches: list, sent_matches: list,
                    attachments: list, has_doc_evidence: bool) -> tuple[SearchStatus, EvidenceStrength]:
    """Classify the overall search outcome."""
    total_msgs = len(outbox_matches) + len(sent_matches)

    if has_doc_evidence:
        return SearchStatus.ATTACHMENT_EVIDENCE_FOUND, EvidenceStrength.STRONG

    if sent_matches and any(
        a.get("parse_status", "").startswith("parsed") and a.get("term_matches", {}).get("srv")
        for a in attachments
    ):
        return SearchStatus.SENT_ITEMS_EVIDENCE_FOUND, EvidenceStrength.STRONG

    if attachments:
        saved = [a for a in attachments if a.get("saved")]
        if saved:
            return SearchStatus.ATTACHMENT_METADATA_FOUND, EvidenceStrength.MODERATE

    if outbox_matches and any(
        "SRV-" in m.get("subject", "") or "REQ" in m.get("subject", "")
        for m in outbox_matches
    ):
        return SearchStatus.SUBJECT_MATCH_ONLY, EvidenceStrength.WEAK

    if outbox_matches or sent_matches:
        return SearchStatus.BODY_CONTEXT_FOUND, EvidenceStrength.WEAK

    return SearchStatus.NO_MATCH, EvidenceStrength.NONE


# ── Local model input builder ──────────────────────────────────────────────
def build_model_input(
    run_id: str,
    card_title: str,
    card_project_id: str,
    search_results: dict,
    attachments: list[dict],
    status: str,
    strength: str,
) -> dict:
    """Build the structured input for the local model."""
    # Summarise evidence
    evidence_items = []
    for m in search_results.get("inbox_matches", []):
        terms = search_text_for_terms(f"{m.get('subject', '')} {m.get('body_preview', '')}")
        evidence_items.append({
            "type": "inbox_email",
            "subject": m.get("subject", "")[:100],
            "sender": m.get("sender", ""),
            "date": m.get("date", ""),
            "source_folder": m.get("source_folder", ""),
            "conversation_id": m.get("conversation_id", ""),
            "term_matches": {k: v[:5] for k, v in terms.items() if v},
            "has_attachments": False,
        })

    for m in search_results.get("sent_matches", []):
        terms = search_text_for_terms(f"{m.get('subject', '')} {m.get('body_preview', '')}")
        evidence_items.append({
            "type": "sent_email",
            "subject": m.get("subject", "")[:100],
            "sender": m.get("sender", ""),
            "date": m.get("date", ""),
            "source_folder": m.get("source_folder", ""),
            "conversation_id": m.get("conversation_id", ""),
            "term_matches": {k: v[:5] for k, v in terms.items() if v},
            "has_attachments": len([a for a in attachments if a.get("email_key") == m.get("email_key")]) > 0,
        })

    att_summaries = []
    for a in attachments:
        tm = a.get("term_matches", {})
        att_summaries.append({
            "filename": a.get("original_filename", ""),
            "ext": a.get("extension", ""),
            "size": a.get("size", 0),
            "sha256": a.get("sha256", "")[:16],
            "parse_status": a.get("parse_status", ""),
            "srv_found": bool(tm.get("srv", [])),
            "req_found": bool(tm.get("req", [])),
            "ips_found": bool(tm.get("ips", [])),
            "people_found": bool(tm.get("people", [])),
            "all_terms": tm.get("all", [])[:20],
            "text_preview": (a.get("text_content", "") or "")[:500],
        })

    return {
        "run_id": run_id,
        "card_id": card_project_id,
        "card_title": card_title,
        "search_status": status,
        "evidence_strength": strength,
        "evidence_count": len(evidence_items),
        "attachment_count": len(attachments),
        "evidence_summary": {
            "inbox_messages": len(search_results.get("inbox_matches", [])),
            "sent_messages": len(search_results.get("sent_matches", [])),
            "attachments": len(attachments),
            "conversation_ids": search_results.get("conv_ids", []),
        },
        "evidence_items": evidence_items,
        "attachment_evidence": att_summaries,
        "search_metadata": {
            "inbox_folder": search_results.get("inbox_folder", ""),
            "sent_folder": search_results.get("sent_folder", ""),
            "body_keyword_search": search_results.get("body_keyword_used", False),
        },
    }


# ── Runner ─────────────────────────────────────────────────────────────────
def run_evidence_pipeline(
    settings_path: str | Path,
    evidence_root: str | Path,
    card_title: str,
    card_project_id: str,
    subject_patterns: list[str] | None = None,
    body_keywords: list[str] | None = None,
    attachment_keywords: list[str] | None = None,
    max_body_search: int = 10,
) -> dict:
    """Run the full evidence pipeline for a card.

    Args:
        settings_path: Path to settings.json.
        evidence_root: Root for evidence output.
        card_title: Human card title.
        card_project_id: Kanban project ID.
        subject_patterns: Subject patterns for DASL search.
        body_keywords: Keywords for body fallback search.
        attachment_keywords: Keywords for attachment filename search.
        max_body_search: Max results from body keyword search.

    Returns:
        Dict with run_id, manifest, model_input, attachments, search_results.
    """
    settings_path = Path(settings_path)
    evidence_root = Path(evidence_root)
    run_id = f"ep_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    run_dir = evidence_root / run_id
    att_dir = run_dir / "extracted_attachments"
    att_dir.mkdir(parents=True, exist_ok=True)

    result = {
        "run_id": run_id,
        "run_dir": str(run_dir),
        "card_title": card_title,
        "card_project_id": card_project_id,
        "search_results": {},
        "attachments": [],
        "manifest": {},
        "model_input": {},
    }

    # Config
    settings = _load_settings(settings_path)
    old_state = _enable_mailbox(settings, settings_path)
    settings = _load_settings(settings_path)

    # Kanban hash before
    kanban_hash_before = ""
    try:
        from .kanban_reader import file_hash, find_projects_json
        from .config import ConfigLoader
        cl = ConfigLoader(str(settings_path) if settings_path.exists() else None)
        ks = cl.load()
        pj, _ = find_projects_json(ks.kanban_local_path())
        if pj and pj.exists():
            kanban_hash_before = file_hash(pj)
    except Exception:
        pass

    # Detect tesseract
    tess_info = detect_tesseract()

    try:
        # Connect Outlook
        app, ns, inbox, sent = _connect_outlook()
        inbox_info = _get_folder_info(inbox)
        sent_info = _get_folder_info(sent)

        search_res = {
            "inbox_matches": [],
            "sent_matches": [],
            "conv_ids": [],
            "body_keyword_used": False,
            "att_search_used": False,
            "inbox_scanned": inbox_info.get("totalItemCount", 0),
            "sent_scanned": sent_info.get("totalItemCount", 0),
            "inbox_folder": inbox_info.get("folderPath", ""),
            "sent_folder": sent_info.get("folderPath", ""),
        }

        all_messages = []
        seen_keys = set()

        # Layer 1: Subject search (Inbox)
        if subject_patterns:
            print(f"  Layer 1: Subject search (Inbox) — {len(subject_patterns)} patterns")
            inbox_matches = search_by_subject(inbox, subject_patterns, "inbox", "Inbox")
            for m in inbox_matches:
                if m["email_key"] not in seen_keys:
                    seen_keys.add(m["email_key"])
                    all_messages.append(m)
            search_res["inbox_matches"] = inbox_matches
            print(f"    → {len(inbox_matches)} inbox matches")

        # Layer 2: Subject search (Sent Items)
        if subject_patterns:
            print(f"  Layer 2: Subject search (Sent Items) — {len(subject_patterns)} patterns")
            sent_matches = search_by_subject(sent, subject_patterns, "sent", "Sent Items")
            for m in sent_matches:
                if m["email_key"] not in seen_keys:
                    seen_keys.add(m["email_key"])
                    all_messages.append(m)
            search_res["sent_matches"] = sent_matches
            print(f"    → {len(sent_matches)} sent matches")

        # Layer 3: ConversationID expansion
        conv_ids = set()
        for m in all_messages:
            cid = m.get("conversation_id", "")
            if cid and len(cid) > 4:
                conv_ids.add(cid)
        search_res["conv_ids"] = list(conv_ids)

        if conv_ids:
            print(f"  Layer 3: ConversationID expansion — {len(conv_ids)} IDs")
            for cid in conv_ids:
                conv_matches = search_by_conversation(inbox, cid, seen_keys, "inbox", "Inbox")
                for m in conv_matches:
                    seen_keys.add(m["email_key"])
                    all_messages.append(m)
                    search_res["inbox_matches"].append(m)
                conv_matches2 = search_by_conversation(sent, cid, seen_keys, "sent", "Sent Items")
                for m in conv_matches2:
                    seen_keys.add(m["email_key"])
                    all_messages.append(m)
                    search_res["sent_matches"].append(m)
            print(f"    → Expanded thread")

        # Layer 4: Body keyword search (Inbox fallback)
        if body_keywords and not search_res["inbox_matches"]:
            print(f"  Layer 4: Body keyword search — {len(body_keywords)} keywords")
            body_matches = search_by_body_keywords(inbox, body_keywords, max_body_search, "inbox", "Inbox")
            for m in body_matches:
                if m["email_key"] not in seen_keys:
                    seen_keys.add(m["email_key"])
                    all_messages.append(m)
                    search_res["inbox_matches"].append(m)
            search_res["body_keyword_used"] = True
            print(f"    → {len(body_matches)} body matches")

        # Layer 5: Attachment filename search
        if attachment_keywords:
            print(f"  Layer 5: Attachment filename search")
            for akw in attachment_keywords:
                att_matches = search_attachments_in_folder(inbox, akw, "inbox", "Inbox")
                for m in att_matches:
                    if m["email_key"] not in seen_keys:
                        seen_keys.add(m["email_key"])
                        all_messages.append(m)
                        search_res["inbox_matches"].append(m)
            search_res["att_search_used"] = True

        # Extract attachments from all matched messages
        print(f"\n  Extracting attachments from {len(all_messages)} messages...")
        all_atts = []
        for msg in all_messages:
            msg_atts = extract_attachments(msg, att_dir)
            for a in msg_atts:
                a["email_subject"] = msg.get("subject", "")
                a["sender"] = msg.get("sender", "")
                a["date"] = msg.get("date", "")
            all_atts.extend(msg_atts)

        # Parse attachments
        print(f"  Parsing {len(all_atts)} attachments...")
        for a in all_atts:
            if a.get("saved"):
                a = parse_attachment(a, tess_info)

        result["attachments"] = all_atts
        result["search_results"] = search_res

        # Classify
        has_doc_evidence = any(
            a.get("parse_status", "") == "parsed" and
            a.get("term_matches", {}).get("srv", [])
            for a in all_atts
        )
        status, strength = classify_search(
            search_res.get("inbox_matches", []),
            search_res.get("sent_matches", []),
            all_atts,
            has_doc_evidence,
        )

        # Build manifest
        manifest = build_evidence_manifest(
            run_id=run_id,
            search_results=search_res,
            attachments=all_atts,
            inbox_info=inbox_info,
            sent_info=sent_info,
            settings=settings,
        )
        manifest["search_results"]["classified_status"] = status.value
        manifest["search_results"]["evidence_strength"] = strength.value

        # Build model input
        model_input = build_model_input(
            run_id=run_id,
            card_title=card_title,
            card_project_id=card_project_id,
            search_results=search_res,
            attachments=all_atts,
            status=status.value,
            strength=strength.value,
        )
        result["manifest"] = manifest
        result["model_input"] = model_input
        result["search_status"] = status.value
        result["evidence_strength"] = strength.value

        # Write output files
        for fname, data in [
            ("evidence_manifest.json", manifest),
            ("search_results.json", search_res),
            ("local_model_input.json", model_input),
        ]:
            with open(run_dir / fname, "w", encoding="utf-8") as f:
                json.dump(data, f, indent=2, ensure_ascii=False, default=str)

        # Write attachment index
        att_index = []
        for a in all_atts:
            safe_a = {k: v for k, v in a.items() if k not in ("_item",)}
            att_index.append(safe_a)
        with open(run_dir / "attachment_index.json", "w", encoding="utf-8") as f:
            json.dump(att_index, f, indent=2, ensure_ascii=False, default=str)

        # Verify hash
        kanban_hash_after = ""
        try:
            pj, _ = find_projects_json(ks.kanban_local_path())
            if pj and pj.exists():
                kanban_hash_after = file_hash(pj)
        except Exception:
            pass
        manifest["safety"]["kanban_hash_before"] = kanban_hash_before[:16] if kanban_hash_before else ""
        manifest["safety"]["kanban_hash_after"] = kanban_hash_after[:16] if kanban_hash_after else ""
        manifest["safety"]["hash_unchanged"] = (kanban_hash_before == kanban_hash_after)

        # Re-write manifest with hash info
        with open(run_dir / "evidence_manifest.json", "w", encoding="utf-8") as f:
            json.dump(manifest, f, indent=2, ensure_ascii=False, default=str)

        print(f"\n  ✓ Pipeline complete: {status.value} / {strength.value}")
        print(f"  ✓ Run dir: {run_dir}")

    except Exception as e:
        import traceback
        result["error"] = str(e)
        result["traceback"] = traceback.format_exc()
        print(f"\n  ✗ Pipeline error: {e}")

    finally:
        # Restore mailbox config
        settings = _load_settings(settings_path)
        _restore_mailbox(settings, settings_path, old_state)

    return result


# ── Simplified CLI wrapper ────────────────────────────────────────────────


def run_evidence_search(
    settings: Any,
    card_title: str,
    card_project_id: str,
    subject_patterns: list[str] | None = None,
    body_keywords: list[str] | None = None,
    attachment_keywords: list[str] | None = None,
    max_body_search: int = 10,
    evidence_root: str | Path | None = None,
) -> dict:
    """Simplified evidence search wrapper for CLI use.
    Takes a settings object (not path) and resolves paths automatically.
    """
    if evidence_root is None:
        output_root = getattr(settings, "output_root", str(Path.cwd()))
        evidence_root = Path(output_root).parent / "apply" / "evidence"
    evidence_root = Path(evidence_root)
    # Resolve settings path from settings object
    sp = getattr(settings, "config_path", None) or str(Path.cwd() / "config" / "settings.json")
    return run_evidence_pipeline(
        settings_path=sp,
        evidence_root=evidence_root,
        card_title=card_title,
        card_project_id=card_project_id,
        subject_patterns=subject_patterns,
        body_keywords=body_keywords,
        attachment_keywords=attachment_keywords,
        max_body_search=max_body_search,
    )
